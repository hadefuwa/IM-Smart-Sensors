"""
Generate full-content Word documents from CP0001 and CP0002 worksheets.

Extracts contentHtml from the JS source files, parses it with BeautifulSoup,
converts SVGs to PNG images, and builds Word documents.

Output (docs/ folder):
  CP0001.docx          — student worksheet book (content + unanswered MCQs)
  CP0002.docx          — student worksheet book (content + unanswered MCQs)
  CP0001 Answers.docx  — trainer key (content + MCQs with answers highlighted)
  CP0002 Answers.docx  — trainer key (content + MCQs with answers highlighted)
"""

import os, re, io, json, tempfile, textwrap
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString, Tag

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import subprocess

ROOT    = Path(__file__).parent.parent
SRC_DIR = ROOT / 'src'
OUT_DIR = ROOT / 'docs'
IMG_DIR = ROOT / 'public' / 'assets' / 'img'

# ─────────────────────────────────────────────────────────────────────────────
# COLOURS
NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
GREEN  = RGBColor(0x16, 0xA3, 0x4A)
ORANGE = RGBColor(0xEA, 0x58, 0x0C)
AMBER  = RGBColor(0xD9, 0x77, 0x06)
RED    = RGBColor(0xDC, 0x26, 0x26)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
LGRAY  = RGBColor(0xF1, 0xF5, 0xF9)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
BLUE   = RGBColor(0x1D, 0x4E, 0xD8)


# ─────────────────────────────────────────────────────────────────────────────
# STEP 1 — Extract contentHtml template literals from JS source files
# ─────────────────────────────────────────────────────────────────────────────

def extract_template_literal(src: str, start: int) -> tuple[str, int]:
    """
    Given position of an opening backtick, return (content, end_index).
    Handles escaped backticks and nested ${...} expressions.
    """
    i = start + 1
    depth = 0        # nesting depth inside ${ }
    str_char = None  # inside a string within ${ }
    out = []

    while i < len(src):
        c = src[i]

        if c == '\\' and i + 1 < len(src):
            out.append(src[i:i+2])
            i += 2
            continue

        if depth == 0:
            if c == '`':
                return ''.join(out), i
            if c == '$' and i + 1 < len(src) and src[i+1] == '{':
                depth += 1
                out.append('')   # skip interpolation — just drop the ${ ... }
                i += 2
                continue
        else:
            # Inside ${ ... } — skip content until matching }
            if str_char:
                if c == str_char:
                    str_char = None
                i += 1
                continue
            if c in ('"', "'", '`'):
                str_char = c
                i += 1
                continue
            if c == '{':
                depth += 1
            elif c == '}':
                depth -= 1
                if depth == 0:
                    i += 1
                    continue
            i += 1
            continue

        out.append(c)
        i += 1

    return ''.join(out), -1


def extract_worksheets_from_js(js_path: Path) -> list[dict]:
    """Parse a worksheet JS file and return list of worksheet dicts."""
    src = js_path.read_text(encoding='utf-8')
    worksheets = []

    # Find each worksheet object by looking for `id: N,` followed by `title:`
    # then find the contentHtml template literal
    ws_start_pattern = re.compile(r'\{\s*\n?\s*id:\s*(\d+)\s*,')

    for m in ws_start_pattern.finditer(src):
        ws_id = int(m.group(1))
        chunk_start = m.start()

        # Find title
        title_m = re.search(r"title:\s*'([^']*)'", src[chunk_start:chunk_start+400])
        title = title_m.group(1) if title_m else f'Worksheet {ws_id}'

        # Find shortDesc
        desc_m = re.search(r"shortDesc:\s*'([^']*)'", src[chunk_start:chunk_start+400])
        short_desc = desc_m.group(1) if desc_m else ''

        # Find estimatedTime
        time_m = re.search(r"estimatedTime:\s*'([^']*)'", src[chunk_start:chunk_start+600])
        est_time = time_m.group(1) if time_m else ''

        # Find whyItMatters (may span multiple lines or use concatenation)
        why_m = re.search(r"whyItMatters:\s*'([^']*)'", src[chunk_start:chunk_start+800])
        why = why_m.group(1) if why_m else ''

        # Find contentHtml template literal
        content_key = src.find('contentHtml:', chunk_start)
        if content_key == -1:
            continue
        backtick_pos = src.find('`', content_key + len('contentHtml:'))
        if backtick_pos == -1:
            continue

        html_content, end_pos = extract_template_literal(src, backtick_pos)

        worksheets.append({
            'id':           ws_id,
            'title':        title.replace("\\'", "'"),
            'shortDesc':    short_desc,
            'estimatedTime': est_time,
            'whyItMatters': why,
            'contentHtml':  html_content,
        })

    # Deduplicate (same id might match twice if the regex overfires)
    seen = set()
    unique = []
    for ws in worksheets:
        if ws['id'] not in seen:
            seen.add(ws['id'])
            unique.append(ws)
    unique.sort(key=lambda w: w['id'])
    return unique


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2 — SVG → PNG helper
# ─────────────────────────────────────────────────────────────────────────────

SVG_HELPER = Path(__file__).parent / 'svg2png.mjs'

def svg_to_png_bytes(svg_str: str, width_px: int = 560) -> bytes | None:
    """Render an SVG string to PNG bytes using Node.js + sharp."""
    try:
        result = subprocess.run(
            ['node', str(SVG_HELPER), str(width_px)],
            input=svg_str.encode('utf-8'),
            capture_output=True,
            timeout=15,
        )
        if result.returncode == 0 and result.stdout:
            return result.stdout
        if result.stderr:
            print(f'  SVG render warning: {result.stderr.decode()[:120]}')
        return None
    except Exception as e:
        print(f'  SVG render failed: {e}')
        return None


# ─────────────────────────────────────────────────────────────────────────────
# STEP 3 — Word document helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def para_fmt(p, space_before=0, space_after=4, left_indent=None, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)
    if keep:
        pf.keep_with_next = True


def add_run(p, text: str, bold=False, italic=False, color=None, size=None, mono=False):
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    if mono:
        r.font.name = 'Courier New'
    return r


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4 — HTML → Word converter
# ─────────────────────────────────────────────────────────────────────────────

# Elements that contain meaningful content we should skip rendering
# (interactive JS-only elements)
SKIP_TAGS = {'script', 'style', 'button', 'input', 'select', 'option', 'noscript'}
# Input types that DO produce useful text when we handle them specially
CHECKBOX_TYPES = {'checkbox'}
RADIO_TYPES    = {'radio'}

# Classes that indicate an info/callout box
INFO_BOX_CLASSES = {'rounded-xl', 'rounded-lg', 'rounded-md'}
CODE_CLASSES     = {'font-mono', 'code-block', 'mockup-code'}


def get_classes(el) -> set:
    return set((el.get('class') or []))


def classify_div(el) -> str:
    """Return semantic type for a div: 'info', 'code', 'grid', 'plain'."""
    classes = get_classes(el)
    if classes & {'font-mono'} or el.find('code'):
        return 'code'
    if classes & INFO_BOX_CLASSES:
        return 'info'
    if 'grid' in classes:
        return 'grid'
    return 'plain'


def inner_text(el) -> str:
    """Get all text from an element, stripping extra whitespace."""
    return re.sub(r'\s+', ' ', el.get_text(' ', strip=True)).strip()


def process_inline(p, node, bold=False, italic=False, color=None, mono=False):
    """Render inline content (text + inline elements) into paragraph p."""
    if isinstance(node, NavigableString):
        txt = str(node)
        txt = re.sub(r'\s+', ' ', txt)
        if txt.strip():
            add_run(p, txt, bold=bold, italic=italic, color=color, mono=mono)
        elif ' ' in txt:
            add_run(p, ' ', bold=bold, italic=italic, color=color, mono=mono)
        return

    tag = node.name if isinstance(node, Tag) else None
    if tag is None:
        return

    if tag in {'strong', 'b'}:
        for child in node.children:
            process_inline(p, child, bold=True, italic=italic, color=color, mono=mono)
    elif tag in {'em', 'i'}:
        for child in node.children:
            process_inline(p, child, bold=bold, italic=True, color=color, mono=mono)
    elif tag in {'code', 'kbd', 'samp'}:
        for child in node.children:
            process_inline(p, child, bold=bold, italic=italic, color=color, mono=True)
    elif tag == 'span':
        cl = get_classes(node)
        c = color
        if 'text-success' in cl or 'text-green' in cl:
            c = GREEN
        elif 'text-error' in cl or 'text-red' in cl:
            c = RED
        elif 'text-warning' in cl or 'text-amber' in cl:
            c = AMBER
        elif 'text-primary' in cl or 'text-blue' in cl:
            c = BLUE
        for child in node.children:
            process_inline(p, child, bold=bold, italic=italic, color=c, mono=mono)
    elif tag == 'a':
        for child in node.children:
            process_inline(p, child, bold=bold, italic=italic, color=BLUE, mono=mono)
    elif tag == 'br':
        add_run(p, '\n')
    elif tag in {'div', 'p', 'li', 'td', 'th', 'h1', 'h2', 'h3', 'h4', 'label'}:
        # block-level inside inline context — just recurse as inline
        for child in node.children:
            process_inline(p, child, bold=bold, italic=italic, color=color, mono=mono)
    else:
        for child in node.children:
            process_inline(p, child, bold=bold, italic=italic, color=color, mono=mono)


def add_svg_image(doc, svg_str: str):
    """Convert SVG and insert as image paragraph."""
    png = svg_to_png_bytes(svg_str, width_px=520)
    if png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_fmt(p, space_before=4, space_after=4)
        run = p.add_run()
        run.add_picture(io.BytesIO(png), width=Inches(5.5))
    else:
        p = doc.add_paragraph('[Diagram — see interactive worksheet]')
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = GRAY


def process_table_el(doc, table_el):
    """Render an HTML <table> element into a Word table."""
    rows_html = table_el.find_all('tr')
    if not rows_html:
        return
    max_cols = max(len(r.find_all(['td','th'])) for r in rows_html)
    if max_cols == 0:
        return

    t = doc.add_table(rows=len(rows_html), cols=max_cols)
    t.style = 'Table Grid'

    for ri, row_el in enumerate(rows_html):
        cells_html = row_el.find_all(['td','th'])
        for ci, cell_el in enumerate(cells_html):
            if ci >= max_cols:
                break
            cell = t.rows[ri].cells[ci]
            txt = inner_text(cell_el)
            cell.text = txt
            r = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
            is_header = (cell_el.name == 'th') or (ri == 0)
            if is_header:
                set_cell_bg(cell, '1E3A5F')
                if r:
                    r.bold = True
                    r.font.color.rgb = WHITE
            elif ri % 2 == 1:
                set_cell_bg(cell, 'F1F5F9')
    doc.add_paragraph()


def process_node(doc, node, indent=0, in_list=False):
    """Recursively process a BS4 node and add content to doc."""
    if isinstance(node, NavigableString):
        txt = str(node).strip()
        if txt:
            p = doc.add_paragraph(txt)
            para_fmt(p, left_indent=indent)
        return

    if not isinstance(node, Tag):
        return

    tag = node.name

    # ── Skip non-content elements ─────────────────────────────────────────
    if tag in SKIP_TAGS:
        return
    if tag == 'input':
        itype = (node.get('type') or '').lower()
        if itype == 'checkbox':
            return  # handled by parent label
        return
    if tag == 'canvas':
        p = doc.add_paragraph()
        para_fmt(p, left_indent=indent + 0.4)
        add_run(p, '[Interactive animation — see the live worksheet]', italic=True, color=GRAY)
        return

    # ── SVG ───────────────────────────────────────────────────────────────
    if tag == 'svg':
        add_svg_image(doc, str(node))
        return

    # ── Headings ──────────────────────────────────────────────────────────
    if tag in ('h1','h2','h3','h4','h5','h6'):
        level = int(tag[1])
        txt = inner_text(node)
        if not txt:
            return
        hp = doc.add_heading(txt, level=min(level+1, 4))
        hp.runs[0].font.color.rgb = NAVY
        para_fmt(hp, space_before=8, space_after=2)
        return

    # ── Paragraph ─────────────────────────────────────────────────────────
    if tag == 'p':
        txt = inner_text(node)
        if not txt:
            return
        p = doc.add_paragraph()
        para_fmt(p, left_indent=indent)
        process_inline(p, node)
        return

    # ── Code block ────────────────────────────────────────────────────────
    if tag in ('pre', 'code') and node.parent and node.parent.name == 'pre':
        txt = node.get_text()
        lines = txt.split('\n')
        for line in lines:
            p = doc.add_paragraph(line)
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(8.5)
            p.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            para_fmt(p, space_before=0, space_after=0, left_indent=indent + 0.4)
        return

    # ── Table ────────────────────────────────────────────────────────────
    if tag == 'table':
        process_table_el(doc, node)
        return

    # ── Lists ─────────────────────────────────────────────────────────────
    if tag in ('ul', 'ol'):
        for i, child in enumerate(node.find_all('li', recursive=False)):
            txt = inner_text(child)
            if not txt:
                continue
            p = doc.add_paragraph(style='List Bullet' if tag == 'ul' else 'List Number')
            para_fmt(p, left_indent=indent + 0.4, space_before=0, space_after=2)
            process_inline(p, child)
        return

    # ── Label (checklist item) ────────────────────────────────────────────
    if tag == 'label':
        cl = get_classes(node)
        # Skip MCQ radio labels — we handle those from our MCQ data
        radio = node.find('input', {'type': 'radio'})
        if radio:
            return
        # Checkbox-style checklist
        checkbox = node.find('input', {'type': 'checkbox'})
        span = node.find('span') or node
        txt = inner_text(span)
        if not txt:
            return
        p = doc.add_paragraph()
        para_fmt(p, left_indent=indent + 0.4, space_before=1, space_after=1)
        add_run(p, '☐  ')
        process_inline(p, span if span != node else node)
        return

    # ── Divider ───────────────────────────────────────────────────────────
    if tag == 'hr' or ('divider' in get_classes(node) and tag == 'div' and not node.get_text(strip=True)):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        return

    # ── Details / Summary (expandable) ────────────────────────────────────
    if tag == 'details':
        summary = node.find('summary')
        if summary:
            p = doc.add_paragraph()
            para_fmt(p, space_before=4)
            add_run(p, inner_text(summary), bold=True, color=NAVY)
        for child in node.children:
            if isinstance(child, Tag) and child.name == 'summary':
                continue
            process_node(doc, child, indent=indent + 0.4)
        return

    # ── Info / callout box (div with border) ──────────────────────────────
    if tag == 'div':
        cl = get_classes(node)

        # Code-style block (font-mono or leading-none overflow-x-auto)
        is_code = bool(cl & {'font-mono'}) or bool(node.find('code'))
        if is_code:
            txt = node.get_text('\n', strip=True)
            for line in txt.split('\n'):
                if not line.strip():
                    continue
                p = doc.add_paragraph(line)
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(8.5)
                p.runs[0].font.color.rgb = BLUE
                para_fmt(p, space_before=0, space_after=0, left_indent=indent + 0.4)
            doc.add_paragraph()
            return

        # Info/callout box — render children with slight indent
        is_box = bool(cl & INFO_BOX_CLASSES)
        new_indent = indent + (0.4 if is_box else 0)

        # Check if this div contains MCQ radio buttons — skip it (handled by MCQ data)
        if node.find('input', {'type': 'radio'}):
            return

        for child in node.children:
            process_node(doc, child, indent=new_indent)
        return

    # ── Anything else — just recurse ──────────────────────────────────────
    for child in node.children:
        process_node(doc, child, indent=indent)


def process_content_html(doc, html: str):
    """Parse and render a worksheet's contentHtml into the document."""
    soup = BeautifulSoup(html, 'html.parser')
    for child in soup.children:
        process_node(doc, child)


# ─────────────────────────────────────────────────────────────────────────────
# STEP 5 — MCQ data (same as generate_docx.py but keyed by ws id + course)
# ─────────────────────────────────────────────────────────────────────────────

# Imported from the existing generate_docx.py data structure
# to avoid duplication, we import it:
import importlib.util, sys

def load_mcq_data():
    spec = importlib.util.spec_from_file_location(
        'gen_docx', Path(__file__).parent / 'generate_docx.py'
    )
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod.CP0001, mod.CP0002


# ─────────────────────────────────────────────────────────────────────────────
# STEP 6 — Word document page setup helpers
# ─────────────────────────────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin   = Cm(2.5)
        s.right_margin  = Cm(2.5)
    return doc


def add_cover(doc, course_name, subtitle):
    """Add a title page."""
    doc.add_paragraph()
    logo = IMG_DIR / 'matrix.png'
    if not logo.exists():
        logo = ROOT / 'public' / 'matrix.png'
    if logo.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(str(logo), width=Inches(2.0))

    title_p = doc.add_heading(course_name, 0)
    title_p.runs[0].font.color.rgb = NAVY
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_fmt(title_p, space_before=12)

    sub_p = doc.add_paragraph(subtitle)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.runs[0].font.color.rgb = GRAY
    sub_p.runs[0].font.size = Pt(13)
    sub_p.runs[0].italic = True

    doc.add_page_break()


def add_ws_header(doc, ws_meta: dict):
    """Add worksheet title, time estimate, and why-it-matters block."""
    h = doc.add_heading(f"WS{ws_meta['id']} — {ws_meta['title']}", level=1)
    h.runs[0].font.color.rgb = NAVY
    para_fmt(h, space_before=6, space_after=4)

    if ws_meta.get('estimatedTime'):
        p = doc.add_paragraph()
        add_run(p, f"⏱  {ws_meta['estimatedTime']}", color=GRAY, size=9.5)
        para_fmt(p, space_before=0, space_after=2)

    if ws_meta.get('whyItMatters'):
        p = doc.add_paragraph()
        add_run(p, 'Why it matters:  ', bold=True, color=NAVY, size=9.5)
        add_run(p, ws_meta['whyItMatters'], italic=True, color=GRAY, size=9.5)
        para_fmt(p, space_before=0, space_after=6)


def add_mcq_section(doc, mcq_ws: dict, show_answers: bool):
    """Add the Knowledge Check MCQ section from our MCQ data dict."""
    if not mcq_ws.get('questions') and not mcq_ws.get('scenarios') and not mcq_ws.get('freetext') and not mcq_ws.get('tables'):
        return

    h = doc.add_heading('Knowledge Check', level=2)
    h.runs[0].font.color.rgb = NAVY
    para_fmt(h, space_before=8, space_after=4)

    # Tables (matching/classification)
    for tbl in mcq_ws.get('tables', []):
        hp = doc.add_paragraph()
        para_fmt(hp, space_before=4, space_after=2)
        add_run(hp, tbl['header'], bold=True, color=NAVY)

        cols = tbl['cols']
        rows = tbl['answers'] if show_answers else tbl['rows']
        t = doc.add_table(rows=1 + len(rows), cols=len(cols))
        t.style = 'Table Grid'
        hdr = t.rows[0]
        for ci, cn in enumerate(cols):
            hdr.cells[ci].text = cn
            hdr.cells[ci].paragraphs[0].runs[0].bold = True
            hdr.cells[ci].paragraphs[0].runs[0].font.color.rgb = WHITE
            set_cell_bg(hdr.cells[ci], '1E3A5F')
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri+1].cells[ci]
                cell.text = str(val)
                if show_answers and ci == len(cols)-1:
                    r = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                    if r:
                        r.bold = True
                        r.font.color.rgb = GREEN
                if ri % 2 == 1:
                    set_cell_bg(cell, 'F1F5F9')
        doc.add_paragraph()

    # Free-text tasks
    for ft in mcq_ws.get('freetext', []):
        p = doc.add_paragraph()
        para_fmt(p, space_before=6, space_after=2)
        add_run(p, ft['q'], bold=True)
        if show_answers:
            ap = doc.add_paragraph()
            para_fmt(ap, left_indent=0.8, space_before=0, space_after=4)
            add_run(ap, f"Answer: {ft['answer']}", bold=True, color=GREEN)
        else:
            doc.add_paragraph()
            doc.add_paragraph()

    # MCQs
    for qi, q in enumerate(mcq_ws.get('questions', []), 1):
        qp = doc.add_paragraph()
        para_fmt(qp, space_before=6, space_after=2)
        add_run(qp, f'Q{qi}.  ', bold=True)
        add_run(qp, q['q'])

        for letter, text in q['options']:
            is_correct = (letter == q['correct'])
            op = doc.add_paragraph()
            para_fmt(op, left_indent=0.8, space_before=1, space_after=1)
            if show_answers and is_correct:
                add_run(op, '✓  ', bold=True, color=GREEN)
                add_run(op, f'{letter})  ', bold=True, color=GREEN)
                add_run(op, text, bold=True, color=GREEN)
            else:
                add_run(op, f'     {letter})  ', color=GRAY)
                add_run(op, text)

        if show_answers and q.get('note'):
            np = doc.add_paragraph()
            para_fmt(np, left_indent=0.8, space_before=1, space_after=2)
            add_run(np, q['note'], italic=True, color=GRAY, size=9)

    # Scenario questions
    for sc in mcq_ws.get('scenarios', []):
        sp = doc.add_paragraph()
        para_fmt(sp, space_before=8, space_after=2)
        add_run(sp, sc['title'], bold=True, color=ORANGE)

        if sc.get('intro'):
            ip = doc.add_paragraph()
            para_fmt(ip, left_indent=0.4, space_before=0, space_after=2)
            add_run(ip, sc['intro'], italic=True, color=GRAY)

        if 'options' not in sc:
            if sc.get('note'):
                np = doc.add_paragraph()
                para_fmt(np, left_indent=0.4)
                add_run(np, sc['note'], italic=True, color=GRAY)
            continue

        for letter, text in sc['options']:
            is_correct = (letter == sc.get('correct'))
            op = doc.add_paragraph()
            para_fmt(op, left_indent=0.8, space_before=1, space_after=1)
            if show_answers and is_correct:
                add_run(op, '✓  ', bold=True, color=GREEN)
                add_run(op, f'{letter})  ', bold=True, color=GREEN)
                add_run(op, text, bold=True, color=GREEN)
            else:
                add_run(op, f'     {letter})  ', color=GRAY)
                add_run(op, text)

        if show_answers and sc.get('note'):
            np = doc.add_paragraph()
            para_fmt(np, left_indent=0.8, space_before=1)
            add_run(np, f'Note: {sc["note"]}', italic=True, color=GRAY, size=9)


# ─────────────────────────────────────────────────────────────────────────────
# STEP 7 — Build one document for a course
# ─────────────────────────────────────────────────────────────────────────────

def build_course_doc(course_name: str, js_path: Path, mcq_list: list, show_answers: bool, subtitle: str) -> Document:
    print(f'\nBuilding: {course_name} {"(Answers)" if show_answers else ""}')
    worksheets = extract_worksheets_from_js(js_path)
    print(f'  Extracted {len(worksheets)} worksheets from {js_path.name}')

    # Build MCQ lookup by ws number (1-based)
    mcq_by_id = {}
    for i, mcq_ws in enumerate(mcq_list, 1):
        mcq_by_id[i] = mcq_ws

    doc = new_doc()
    add_cover(doc, course_name, subtitle)

    for ws in worksheets:
        ws_id = ws['id']
        print(f'  WS{ws_id}: {ws["title"]}  ({len(ws["contentHtml"])} chars of HTML)')

        add_ws_header(doc, ws)
        process_content_html(doc, ws['contentHtml'])

        # Append MCQ section if we have MCQ data for this worksheet
        if ws_id in mcq_by_id:
            add_mcq_section(doc, mcq_by_id[ws_id], show_answers=show_answers)

        doc.add_page_break()

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    CP0001_MCQ, CP0002_MCQ = load_mcq_data()

    tasks = [
        ('CP0001.docx',         'CP0001 — Interactive Worksheets', SRC_DIR/'worksheets-page.js', CP0001_MCQ, False, 'Student Workbook'),
        ('CP0001 Answers.docx', 'CP0001 — Interactive Worksheets', SRC_DIR/'worksheets-page.js', CP0001_MCQ, True,  'Trainer Answer Key'),
        ('CP0002.docx',         'CP0002 — Engineering Worksheets',  SRC_DIR/'cp0002-page.js',    CP0002_MCQ, False, 'Student Workbook'),
        ('CP0002 Answers.docx', 'CP0002 — Engineering Worksheets',  SRC_DIR/'cp0002-page.js',    CP0002_MCQ, True,  'Trainer Answer Key'),
    ]

    for filename, course, js_path, mcq, answers, subtitle in tasks:
        out_path = OUT_DIR / filename
        doc = build_course_doc(course, js_path, mcq, answers, subtitle)
        doc.save(out_path)
        size_kb = out_path.stat().st_size // 1024
        print(f'  Saved: {out_path.name}  ({size_kb} KB)')

    print('\nDone.')
