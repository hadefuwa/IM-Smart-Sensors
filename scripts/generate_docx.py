"""
Generate CP0001.docx, CP0002.docx, CP0001 Answers.docx, CP0002 Answers.docx
from the MCQ data defined in this script.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'docs')

# ─────────────────────────────────────────────────────────────────────────────
# DATA  –  Each worksheet is a dict with:
#   title   : str
#   intro   : str (optional brief context shown before questions)
#   questions: list of dicts, each with:
#       q        : question text
#       options  : list of (letter, text)  e.g. [('A','...'), ('B','...'), ('C','...')]
#       correct  : letter string e.g. 'A'
#       note     : optional str shown only in answer doc (source / explanation)
#   tables   : optional list of matching/decode tables shown in both docs
#              each table is {'header': str, 'rows': [(col1,col2,...)], 'cols':['Col1','Col2',...], 'answer_col': int}
#   scenarios: optional list of scenario MCQs (shown after MCQs)
#   freetext : optional list of {'q':str, 'answer':str} shown as short-answer items
# ─────────────────────────────────────────────────────────────────────────────

CP0001 = [
    {
        'title': 'WS1 — Meet Your Kit',
        'questions': [
            {
                'q': 'What colour is the IO-Link Master, and what is it sometimes called?',
                'options': [('A','Orange — it is called the hub'),
                            ('B','Blue — it is called the controller'),
                            ('C','Green — it is called the gateway')],
                'correct': 'A',
                'note': 'Taught by: the "Meet the Kit" checklist in WS1.',
            },
            {
                'q': 'Which sensor detects metal objects using an electromagnetic field?',
                'options': [('A','The capacitive sensor on Port 2'),
                            ('B','The temperature sensor on Port 3'),
                            ('C','The proximity sensor on Port 1')],
                'correct': 'C',
                'note': 'Taught by: WS1 checklist — "Proximity Sensor… Creates an electromagnetic field".',
            },
            {
                'q': 'What can the capacitive sensor detect that the proximity sensor cannot?',
                'options': [('A','Objects moving faster than 1 m/s'),
                            ('B','Materials like liquid or powder — even through a container wall'),
                            ('C','The exact temperature of an object')],
                'correct': 'B',
                'note': 'Taught by: WS1 checklist — "Capacitive Sensor — Detects materials like liquid, powder, or plastic — even through a container wall".',
            },
            {
                'q': "What is the Raspberry Pi's job in this system?",
                'options': [('A','It collects the sensor data, runs the comms, and serves up this dashboard'),
                            ('B','It powers the IO-Link Master via USB'),
                            ('C','It directly controls the light stack colours')],
                'correct': 'A',
                'note': 'Taught by: WS1 checklist — Edge Device item.',
            },
            {
                'q': 'During the challenge, what did you notice about the capacitive sensor?',
                'options': [('A','It only ever shows two states — on or off — same as a normal sensor'),
                            ('B','It measured the exact distance between your hand and the sensor in millimetres'),
                            ('C','It showed a live level rising as your hand got close, before it fully triggered')],
                'correct': 'C',
                'note': 'Demonstrated by: the live capacitive challenge interaction in WS1.',
            },
        ],
    },
    {
        'title': 'WS2 — What is a Smart Sensor?',
        'questions': [
            {
                'q': 'A basic digital (on/off) proximity sensor fails and its output goes silent. What diagnostic information can a technician retrieve from it?',
                'options': [('A','Nothing — the sensor just goes silent, you have to go and test it manually'),
                            ('B','The exact fault code and which internal component has failed'),
                            ('C','The last reading before it failed, stored in its internal log')],
                'correct': 'A',
                'note': 'Taught by: the Digital Sensor section in WS2.',
            },
            {
                'q': 'A pressure transmitter outputs 12 mA on a 4–20 mA loop calibrated 0–100 bar. What is the pressure?',
                'options': [('A','12 bar — the mA value equals the bar reading directly'),
                            ('B','0 bar — the sensor is in fault because 12 mA is below the normal 16 mA minimum'),
                            ('C','50 bar — 12 mA is the midpoint of the 4–20 mA range, which maps to 50% of 100 bar')],
                'correct': 'C',
                'note': 'Taught by: the Analogue Sensor (4–20 mA) section in WS2.',
            },
            {
                'q': 'An IO-Link sensor still has OUT1 active. What does the IO-Link channel add on top?',
                'options': [('A','Process data, fault codes, device identity, and remote parameter writes — all over the same 3-wire cable without interrupting OUT1'),
                            ('B','It replaces OUT1 entirely — you only get the data stream, not a switching output'),
                            ('C','A higher voltage signal so the PLC can tell it is an IO-Link device')],
                'correct': 'A',
                'note': 'Taught by: the IO-Link Smart Sensor bullet points in WS2.',
            },
            {
                'q': 'Why choose an analogue sensor over a digital sensor for tank level monitoring?',
                'options': [('A','Analogue sensors are cheaper and easier to wire than digital sensors'),
                            ('B','Analogue sensors send fault codes that digital sensors cannot'),
                            ('C','An analogue signal gives a continuous level reading — you see exactly how full the tank is, not just "full" or "empty"')],
                'correct': 'C',
                'note': 'Taught by: the Analogue Sensor section in WS2.',
            },
            {
                'q': 'In the IO-Link wire animation, what does the square wave at the bottom of the cable represent?',
                'options': [('A','The IO-Link data packets being encoded onto the wire'),
                            ('B','The standard switching output (OUT1) — still present and working exactly as a normal digital sensor, while the data stream runs above it'),
                            ('C','The power supply waveform showing 24V DC to the sensor')],
                'correct': 'B',
                'note': 'Taught by: the IO-Link animation in WS2.',
            },
        ],
        'scenarios': [
            {
                'title': 'Fault Diagnosis Scenario',
                'intro': 'The scenario injects a randomised fault from 5 possibilities. Read the debug console hints and select the most likely cause.',
                'note': 'Randomised each run — no single correct letter. Correct answer revealed through hints in the debug console.',
            }
        ],
    },
    {
        'title': 'WS3 — The Proximity Sensor',
        'questions': [
            {
                'q': 'What does the OUT1 chart do when metal is detected?',
                'options': [('A','The line jumps from 0 to 1, then drops back to 0 when the metal moves away'),
                            ('B','The line stays flat at 0 the whole time'),
                            ('C','The line drops below zero when metal is detected')],
                'correct': 'A',
                'note': 'Demonstrated by: the live OUT1 chart task in WS3.',
            },
            {
                'q': 'The Instability Alarm fires while OUT1 is ON. What does this tell you?',
                'options': [('A','The sensor has detected two metal objects at the same time'),
                            ('B','The sensor cable has a loose connection'),
                            ('C','The target is near the edge of the sensing range — detection may be unreliable')],
                'correct': 'C',
                'note': 'Taught by: the PDin bit table in WS3.',
            },
            {
                'q': 'The Over-Approach Alarm fires. What is the correct response?',
                'options': [('A','Increase the standoff distance between the sensor face and the target'),
                            ('B','Replace the sensor immediately — the over-approach alarm means it has failed'),
                            ('C','Reduce sensitivity using the teach button')],
                'correct': 'A',
                'note': 'Taught by: the PDin bit table in WS3 — "Over-Approach Alarm: The target is too close… Increase the standoff distance".',
            },
        ],
        'scenarios': [
            {
                'title': 'Maintenance Scenario — ISDU Misconfiguration',
                'intro': 'The scenario injects a fault: the proximity sensor output logic is set to NC (inverted). The output is ON when nothing is detected and OFF when metal is present. Work through the ISDU read steps, then select the root cause.',
                'options': [('1','Output logic parameter changed from NO to NC — inverts the switching behaviour'),
                            ('2','Sensor coil is damaged — mechanical impact during removal reversed the output'),
                            ('3','IO-Link cable polarity swapped at the master connector — wiring error')],
                'correct': '1',
                'note': 'Correct: option 1. JS checks data-ans="nc". Taught by: the ISDU read/write steps in WS3.',
            }
        ],
    },
    {
        'title': 'WS4 — The Capacitive Sensor',
        'questions': [
            {
                'q': 'What triggers a new count?',
                'options': [('A','Every second while your hand is touching'),
                            ('B','Only when you pull your hand away'),
                            ('C','The moment detection goes from off to on — one touch = one count, no matter how long you hold it')],
                'correct': 'C',
                'note': 'Demonstrated by: the live detection counter in WS4.',
            },
            {
                'q': 'Which materials can a capacitive sensor detect that a proximity sensor cannot?',
                'options': [('A','Water and powder — even through a container wall'),
                            ('B','Sound and light'),
                            ('C','Heat and pressure')],
                'correct': 'A',
                'note': 'Taught by: the WS4 introduction and SVG diagram.',
            },
            {
                'q': 'The capacitive sensor output is always ON even when nothing is near it. Most likely cause?',
                'options': [('A','Sensitivity too high — it is detecting the container wall or nearby objects'),
                            ('B','The sensor needs replacing immediately'),
                            ('C','The cable is the wrong colour')],
                'correct': 'A',
                'note': 'Taught by: the Setpoint section in WS4.',
            },
        ],
    },
    {
        'title': 'WS5 — The Temperature Sensor',
        'questions': [
            {
                'q': 'What extra can you do with an IO-Link temperature sensor vs. a basic switch?',
                'options': [('A','Nothing extra — it works exactly the same as a switch'),
                            ('B','It only works at high temperatures above 100 °C'),
                            ('C','See the actual live temperature in °C, trend it over time, and get early warning before the trip point is reached')],
                'correct': 'C',
                'note': 'Taught by: the three-way comparison box and TV7105 info box in WS5.',
            },
            {
                'q': 'Move SP1 just above current temperature and hold the sensor. What should happen?',
                'options': [('A','Nothing changes — the alarm only activates from the main dashboard'),
                            ('B','The alarm state turns red and shows "ABOVE THRESHOLD" once the temperature crosses your slider value'),
                            ('C','The sensor turns itself off to prevent overheating')],
                'correct': 'B',
                'note': 'Demonstrated by: the live setpoint task in WS5.',
            },
            {
                'q': 'Temperature reading suddenly drops to −40 °C in a room-temperature lab. Most likely cause?',
                'options': [('A','Broken or disconnected sensor — −40 °C is the bottom of the TV7105 measurement range and appears when the sensing element is open-circuit'),
                            ('B','The lab is very cold'),
                            ('C','The setpoint has been changed')],
                'correct': 'A',
                'note': 'Taught by: the fault identification section in WS5.',
            },
        ],
        'scenarios': [
            {
                'title': 'Calibration Drift Scenario (Work Order CAL-0189)',
                'intro': 'The scenario injects a +3.0 °C calibration offset. The sensor reads 3 °C higher than the reference thermometer. After confirming via ISDU read, select the correct corrective write.',
                'options': [('1','Write +3.0 °C — add 3 degrees to the existing offset'),
                            ('2','Write 0.0 °C — restore factory default (no offset)'),
                            ('3','Write −6.0 °C — apply a negative 6-degree correction')],
                'correct': '2',
                'note': 'Correct: option 2. Write 0.0°C to index 681, subindex 0 (int16, scale 0.1). Taught by: the calibration scenario explanation in WS5.',
            }
        ],
    },
    {
        'title': 'WS6 — The Light Stack',
        'questions': [
            {
                'q': 'The CL50 is "PDout-only". What does that mean?',
                'options': [('A','The master sends a command value TO the light — no measurement data comes back from the device to the master'),
                            ('B','It only works with certain IO-Link masters'),
                            ('C','It uses a separate digital output wire for each colour segment')],
                'correct': 'A',
                'note': 'Taught by: the WS6 introduction.',
            },
            {
                'q': 'The hex value 180101 is sent to the CL50. Which colour will it show?',
                'options': [('A','Green — Octet2 low nibble = 0x00'),
                            ('B','Amber — 0x18 in Octet0 indicates Amber'),
                            ('C','Red — Octet2 low nibble = 0x01 = Red (index 1)')],
                'correct': 'C',
                'note': 'Taught by: the PDout table and colour index list in WS6. Octet2 = 0x01 → low nibble = 1 = Red.',
            },
            {
                'q': 'Flashing red will not turn off after the fault is cleared. Most likely cause?',
                'options': [('A','The PLC is still sending a "red on" PDout command — the fault state is held in the controller, not the light itself'),
                            ('B','The light stack hardware is broken'),
                            ('C','IO-Link has lost connection to the light stack')],
                'correct': 'A',
                'note': 'Taught by: the WS6 scenario context — "the PLC output register is still holding the pre-maintenance fault PDout value".',
            },
            {
                'q': 'Why is IO-Link PDout better than a traditional 3-wire digital light stack?',
                'options': [('A','It is always cheaper per unit due to IO-Link certification reducing manufacturing costs'),
                            ('B','The colour cannot be changed once wired — it is safer'),
                            ('C','Full control of colour, animation, intensity, and speed over a single 3-pin cable — plus the HMI can read back the exact state without a separate feedback wire')],
                'correct': 'C',
                'note': 'Taught by: the WS6 introduction and the "Build Your Own" section.',
            },
        ],
        'scenarios': [
            {
                'title': 'Maintenance Scenario (Work Order LST-0312)',
                'intro': 'The light is showing Flashing Amber but all sensors and the PLC are healthy. Select the most likely root cause.',
                'options': [('A','The CL50 hardware is malfunctioning — the amber segment is stuck and cannot be cleared by software'),
                            ('B','IO-Link comms to Port 4 have dropped — the light is locked to its last received command and cannot be updated'),
                            ('C','The PLC output register is still holding the pre-maintenance fault PDout value (184203 — Amber Flash). The normal-run value (000100 — Green Steady) was never sent after the PLC restarted')],
                'correct': 'C',
                'note': 'Correct: C. JS checks sel.value === "c". Taught by: Q3 above; this scenario is direct reinforcement.',
            }
        ],
    },
    {
        'title': 'WS7 — Fault Finding and Replacement',
        'intro': 'Recall — Switching Output Logic (ISDU Index 61 / Sub 1)\n\nNO (Normal Open, value 0): Output is ON when an object is detected. Factory default.\nNC (Normal Closed, value 1): Output is ON when no object is detected — inverted behaviour.',
        'scenarios': [
            {
                'title': 'Maintenance Scenario — Port 1 OUT1 Permanently ON',
                'intro': 'Port 1 OUT1 is permanently ON with nothing in front of the sensor. The P1 dot in the HMI is blue — IO-Link is communicating normally. Your manager says he did not touch Port 1. Select the most likely explanation.',
                'options': [('A','IO-Link communication to Port 1 has dropped — the master is frozen on the last reported output state and cannot be updated'),
                            ('B','The inductive sensor hardware has failed — the output transistor is shorted and permanently conducting regardless of whether an object is present'),
                            ('C','The sensor output logic is set to NC (Normal Closed) — in NC mode OUT1 is ON when nothing is detected and turns OFF when an object is present. This is the inverse of the correct NO (Normal Open) setting and causes the controller to see a permanent "object present" signal')],
                'correct': 'C',
                'note': 'Correct: C. JS checks sel.value === "c". Taught by: WS3 NC/NO scenario + the recall callout in WS7.',
            }
        ],
    },
    {
        'title': 'WS8 — Final Practical Assessment',
        'intro': 'Quick Reference (available in worksheet as expandable panel)\n\nISDU Decode Formula: raw decimal × scale factor = engineering value\nExample: 01 F4 hex → 500 decimal → × 0.1 = 50.0 °C\nNote: int16 values above 32767 are negative (two\'s complement).\n\nCL50 Colour Index (Octet2 low nibble):\n0=Green  1=Red  2=Orange  3=Amber  4=Yellow  5=Lime  6=Spring Green\n7=Cyan  8=Sky Blue  9=Blue  10=Violet  11=Magenta  12=Rose  13=White',
        'freetext': [
            {
                'q': 'Task 2A — Hex Decode: Index 583 (SP1) returns 01 F4 (int16, scale ×0.1). What is the setpoint in °C?',
                'answer': '50.0 °C\n(0x01F4 = 500 decimal; 500 × 0.1 = 50.0)',
            },
            {
                'q': 'Task 3 — PDout Colour: PDout value 000109 is sent to the CL50. What colour will it show?',
                'answer': 'Blue\n(Octet2 = 0x09; low nibble = 9 = Blue — see colour index above)',
            },
        ],
        'questions': [
            {
                'q': 'Task 4 / Q1 — The Omron E2E instability alarm fires but OUT1 is still switching correctly. What is the right course of action?',
                'options': [('A','Ignore it — the output is working so there is no problem'),
                            ('B','Replace the sensor immediately'),
                            ('C','Investigate — instability means the target is at the edge of sensing range. Small vibrations will cause intermittent output. Reposition the target or adjust the sensing distance before it causes a fault')],
                'correct': 'C',
                'note': 'Taught by: WS3 PDin bit table.',
            },
            {
                'q': 'Task 4 / Q2 — After replacing the capacitive sensor on Port 2, its output is ON even with nothing near it. Re-teaching did not fix it. What is the most likely cause?',
                'options': [('A','Output logic is set to NC — the sensor output is inverted so it is ON when nothing is present. Read the output logic parameter via ISDU to confirm, then write NO (0) to correct it'),
                            ('B','The IO-Link cable polarity is reversed — swap Pin 2 and Pin 4 at the master connector'),
                            ('C','SP1 is set correctly — the container wall is triggering the output and this is expected behaviour for a capacitive sensor')],
                'correct': 'A',
                'note': 'Taught by: WS3 NC/NO ISDU scenario and WS4 setpoint section.',
            },
            {
                'q': 'Task 4 / Q3 — A temperature sensor reading drifts 3 °C high after months in service. No other system changes were made. What is the correct IO-Link field action?',
                'options': [('A','Replace the sensor — drift always indicates a faulty unit'),
                            ('B','Write a −3.0 °C calibration offset to the sensor via ISDU (index 681) to correct the reading in the field, verified against a reference thermometer'),
                            ('C','Adjust SP1 down by 3 °C to compensate for the drift')],
                'correct': 'B',
                'note': 'Taught by: WS5 calibration drift scenario.',
            },
            {
                'q': 'Task 4 / Q4 — What is the key advantage of the IO-Link light stack over a conventionally wired one?',
                'options': [('A','IO-Link light stacks use less power'),
                            ('B','IO-Link light stacks can display more colours than conventional ones'),
                            ('C','All colours, animation patterns, and intensity levels are set by a single 3-byte PDout value over one standard cable. Conventional wiring needs a separate signal wire per segment')],
                'correct': 'C',
                'note': 'Taught by: WS6 introduction and Build-Your-Own section.',
            },
        ],
    },
]

CP0002 = [
    {
        'title': 'WS1 — Your System: A Technical Brief',
        'questions': [
            {
                'q': 'What is the IP address of the AL1350 IO-Link Master?',
                'options': [('A','192.168.7.4'),
                            ('B','192.168.7.2'),
                            ('C','192.168.1.1')],
                'correct': 'A',
                'note': 'Taught by: the System Architecture checklist in WS1.',
            },
            {
                'q': 'The AL1350 pushes data every 500 ms. How many WebSocket messages per second does the browser receive?',
                'options': [('A','20 messages per second'),
                            ('B','2 messages per second'),
                            ('C','500 messages per second — 500 ms equals 500 per second')],
                'correct': 'B',
                'note': 'Taught by: MQTT Push section. Formula: 1000 ÷ 500 = 2 per second. Option C is the classic confusion between the period (500 ms) and the frequency.',
            },
            {
                'q': 'The Raspberry Pi has eth0 and wlan0. Why does this matter for the kit?',
                'options': [('A','It allows two IO-Link masters to be connected simultaneously'),
                            ('B','It doubles bandwidth by bonding both interfaces for MQTT traffic'),
                            ('C','It bridges the isolated IO-Link subnet to the building LAN')],
                'correct': 'C',
                'note': 'Taught by: the Network Architecture section in WS1.',
            },
            {
                'q': 'How many IO-Link ports have a sensor connected in this kit?',
                'options': [('A','4 — Ports 1 through 4 each have a sensor'),
                            ('B','8 — every port on the AL1350 is populated'),
                            ('C','2 — only the photoelectric and capacitive sensors')],
                'correct': 'A',
                'note': 'Taught by: the Kit Overview in WS1.',
            },
        ],
    },
    {
        'title': 'WS2 — How This App Works',
        'questions': [
            {
                'q': 'Why is MQTT better suited than HTTP polling for 500 ms sensor delivery?',
                'options': [('A','MQTT uses publish/subscribe so the AL1350 pushes data the moment it is ready, without the Pi having to request each cycle'),
                            ('B','HTTP is push-based and uses persistent connections that the broker manages'),
                            ('C','MQTT guarantees delivery of every message by queuing until the subscriber acknowledges receipt')],
                'correct': 'A',
                'note': 'Taught by: the MQTT vs HTTP section in WS2.',
            },
            {
                'q': 'The AL1350 loses its MQTT subscriptions on power-cycle. How does the backend handle this?',
                'options': [('A','It waits for the user to manually re-subscribe'),
                            ('B','It switches permanently to HTTP polling'),
                            ('C','It calls ensure_mqtt_subscription() on every backend startup, re-registering the push subscription automatically')],
                'correct': 'C',
                'note': 'Taught by: the Backend Startup section in WS2.',
            },
            {
                'q': 'If the FastAPI backend crashed, what would you observe on this page?',
                'options': [('A','Live data would stop and the WebSocket would disconnect'),
                            ('B','The WebSocket message counter would continue to increment at the normal rate'),
                            ('C','The AL1350 would automatically restart the backend service via its own built-in watchdog')],
                'correct': 'A',
                'note': 'Taught by: the Data Flow section in WS2.',
            },
            {
                'q': 'What does consistently high round-trip latency on Connection Diagnostics indicate?',
                'options': [('A','The browser\'s chart rendering pipeline cannot keep up with the WebSocket update rate'),
                            ('B','The AL1350 is under load or there is congestion on the IO-Link subnet'),
                            ('C','The Pi\'s Wi-Fi interface is saturated by other traffic on the building network')],
                'correct': 'B',
                'note': 'Taught by: the Connection Diagnostics section in WS2.',
            },
        ],
    },
    {
        'title': 'WS3 — How Each Sensor Works (Protocol Level)',
        'questions': [
            {
                'q': 'Q1 (Port 1 — Proximity): Raw PDin hex 0x01 — what is the switching output state?',
                'options': [('A','Object detected (bit 0 = 1)'),
                            ('B','No object detected'),
                            ('C','Sensor fault')],
                'correct': 'A',
                'note': 'Taught by: the Omron E2E PDin bit table in WS3.',
            },
            {
                'q': 'Q2 (Port 2 — Capacitive): Why is a detection count more useful than just the switching state?',
                'options': [('A','It gives a continuous analogue level reading rather than a binary on/off output'),
                            ('B','It prevents false positives caused by water or foam near the sensor face'),
                            ('C','A running count tracks how many containers have been filled without a separate counter sensor')],
                'correct': 'C',
                'note': 'Taught by: the Capacitive PDin section in WS3.',
            },
            {
                'q': 'Q3 (Port 3 — Temperature): Raw PDin 0x02EE decodes to what temperature?',
                'options': [('A','7.5 °C — wrong scale factor (÷100 instead of ÷10)'),
                            ('B','750.0 °C — no scaling applied (raw decimal value)'),
                            ('C','75.0 °C')],
                'correct': 'C',
                'note': 'Formula: 0x02EE = 750 decimal; 750 ÷ 10 = 75.0 °C. Taught by: the TV7105 PDin section in WS3.',
            },
            {
                'q': 'Q4 (Port 4 — CL50): Why does building a PDout command for the CL50 require more care than reading a proximity sensor\'s PDin?',
                'options': [('A','Multiple fields — colour, animation, intensity, and speed — must be packed into specific bit positions across 3 bytes. One wrong bit changes the entire behaviour'),
                            ('B','It uses a completely different IO-Link variant that the standard binary decoder cannot process'),
                            ('C','The CL50 PDin arrives as a plain text string rather than binary bytes')],
                'correct': 'A',
                'note': 'Taught by: the PDout Structure section in WS3.',
            },
        ],
    },
    {
        'title': 'WS4 — Benefits for Maintenance',
        'tables': [
            {
                'header': 'Matching Exercise — drag each IO-Link feature to the maintenance problem it solves.',
                'cols': ['Maintenance Problem', 'IO-Link Feature (process / event / service)'],
                'rows': [
                    ('Sensor fails without warning', '___________'),
                    ('Replacing a sensor loses its configuration', '___________'),
                    ('Unknown which sensor on the line failed', '___________'),
                    ('Planning when to replace a sensor', '___________'),
                ],
                'answers': [
                    ('Sensor fails without warning', 'event data'),
                    ('Replacing a sensor loses its configuration', 'service data (parameter store)'),
                    ('Unknown which sensor on the line failed', 'event data'),
                    ('Planning when to replace a sensor', 'event data (lifecycle counter)'),
                ],
            }
        ],
        'questions': [
            {
                'q': 'Which data type does the WebSocket push to the browser on every tick?',
                'options': [('A','Process data (sensor values) + event data (faults/warnings) merged into a single JSON payload'),
                            ('B','Service data only (parameters)'),
                            ('C','Raw binary PDin only')],
                'correct': 'A',
                'note': 'Taught by: the WebSocket Data section in WS6 (Diagnostics worksheet).',
            },
        ],
    },
    {
        'title': 'WS5 — Decoding IO-Link Data',
        'tables': [
            {
                'header': 'Calculation Exercises — decode each raw PDin hex value to an engineering value.',
                'cols': ['Hex Value', 'Type', 'Scale Factor', 'Your Answer'],
                'rows': [
                    ('00 E6', 'int16', '×0.1', '___________'),
                    ('00 64', 'uint16', 'none', '___________'),
                    ('01 F4', 'int16', '×0.1', '___________'),
                ],
                'answers': [
                    ('00 E6', 'int16', '×0.1', '23.0 °C  (RP1 — 230 × 0.1)'),
                    ('00 64', 'uint16', 'none', '100 counts  (SP1 capacitive)'),
                    ('01 F4', 'int16', '×0.1', '50.0 °C  (SP1 temperature — 500 × 0.1)'),
                ],
            }
        ],
    },
    {
        'title': 'WS6 — Diagnostics: Process, Service, and Event Data',
        'tables': [
            {
                'header': 'Classification Exercise — label each item as process, event, or service.',
                'cols': ['Data Item', 'Type (process / event / service)'],
                'rows': [
                    ('Temperature value 23.5 °C arriving every 500 ms', '___________'),
                    ('Device replacement detected on Port 2', '___________'),
                    ('Parameter read — filter time = 10 ms', '___________'),
                    ('Object present flag (switching output state)', '___________'),
                    ('Short circuit fault reported', '___________'),
                ],
                'answers': [
                    ('Temperature value 23.5 °C arriving every 500 ms', 'process'),
                    ('Device replacement detected on Port 2', 'event'),
                    ('Parameter read — filter time = 10 ms', 'service'),
                    ('Object present flag (switching output state)', 'process'),
                    ('Short circuit fault reported', 'event'),
                ],
            }
        ],
    },
    {
        'title': 'WS7 — PLC and HMI Integration',
        'questions': [
            {
                'q': 'How does IO-Link process data reach a PLC I/O scan cycle?',
                'options': [('A','The sensor connects directly to a PLC input card; the master acts only as a power supply'),
                            ('B','The sensor transmits binary frames over Wi-Fi directly to the PLC\'s memory address'),
                            ('C','The IO-Link master maps each port\'s PDin into its process image; the PLC reads this via a fieldbus on every scan cycle')],
                'correct': 'C',
                'note': 'Taught by: the PLC Integration section in WS7.',
            },
            {
                'q': 'What is the key advantage of a web HMI over walking to the machine?',
                'options': [('A','Walking to the machine is always preferred because it provides a physical inspection opportunity'),
                            ('B','A web HMI only functions when the machine has a stable internet connection'),
                            ('C','Faster fault identification, multi-machine monitoring, and historical trends without a site visit')],
                'correct': 'C',
                'note': 'Taught by: the HMI Benefits section in WS7.',
            },
            {
                'q': 'What does a spike in the Connection Diagnostics latency graph tell you?',
                'options': [('A','High latency on the IO-Link subnet or the AL1350 being overloaded'),
                            ('B','The dashboard browser tab is consuming too much memory and slowing chart rendering'),
                            ('C','The Pi\'s MQTT broker has stopped forwarding messages and is queueing them internally')],
                'correct': 'A',
                'note': 'Taught by: the Connection Diagnostics section in WS7.',
            },
        ],
        'tables': [
            {
                'header': 'Backend Component Matching — which file handles each function?',
                'cols': ['Function', 'Backend File'],
                'rows': [
                    ('Subscribing to AL1350 MQTT topics', '___________'),
                    ('Decoding raw PDin into °C / object state / CL50 colour', '___________'),
                    ('Pushing decoded data to the browser via WebSocket', '___________'),
                ],
                'answers': [
                    ('Subscribing to AL1350 MQTT topics', 'io_link_fastapi.py'),
                    ('Decoding raw PDin into °C / object state / CL50 colour', 'decoder.py'),
                    ('Pushing decoded data to the browser via WebSocket', 'io_link_fastapi.py'),
                ],
            }
        ],
    },
    {
        'title': 'WS8 — Case Study: Standard vs IO-Link — The Numbers',
        'intro': 'Use the ROI data in the worksheet to answer the following questions.\nStandard approach: 40 faults × 35 min each = 1,400 min = 23.3 h/year\nIO-Link approach: 40 faults × 8 min each = 320 min = 5.3 h/year',
        'questions': [
            {
                'q': 'What is the annual downtime saving when switching to IO-Link?',
                'options': [('A','18 hours per year'),
                            ('B','5.3 hours per year'),
                            ('C','23.3 hours per year')],
                'correct': 'A',
                'note': 'Calculation: 23.3 − 5.3 = 18 h/year. Taught by: the ROI calculation table in WS8.',
            },
            {
                'q': 'At £5,000/hour production value, what is the annual financial saving from IO-Link?',
                'options': [('A','£175,000'),
                            ('B','£26,500'),
                            ('C','£90,000')],
                'correct': 'C',
                'note': 'Calculation: 18 h × £5,000 = £90,000/year. Taught by: the ROI section in WS8.',
            },
            {
                'q': 'Beyond MTTR reduction, which pair of benefits contributes most to lower total cost of ownership?',
                'options': [('A','Built-in wireless backup channel and automatic IP provisioning for all sensors'),
                            ('B','Automatic parameter restore on sensor swap and remote diagnostics via the HMI'),
                            ('C','Lower per-unit sensor cost due to IO-Link certification subsidising manufacturing')],
                'correct': 'B',
                'note': 'Taught by: the TCO Benefits section in WS8.',
            },
            {
                'q': 'Which data type would you export to a CMMS to automate maintenance work orders?',
                'options': [('A','Event data (faults and warnings) — these are the triggers for maintenance actions'),
                            ('B','Only process data (sensor values)'),
                            ('C','Service data only (parameters)')],
                'correct': 'A',
                'note': 'Taught by: the CMMS Integration section in WS8.',
            },
        ],
    },
    {
        'title': 'WS9 — Device Identity: Vendor ID, Device ID & PDin',
        'questions': [
            {
                'q': 'What Vendor ID does Port 1 (Omron proximity) report on the IO-Link Master page?',
                'options': [('A','612'),
                            ('B','310'),
                            ('C','1586')],
                'correct': 'A',
                'note': 'Answer visible live on the Port 1 card; 612 = OMRON Corporation.',
            },
            {
                'q': 'Which manufacturer does Vendor ID 612 identify?',
                'options': [('A','ifm electronic'),
                            ('B','OMRON Corporation'),
                            ('C','Balluff')],
                'correct': 'B',
                'note': 'Taught by: the Vendor Registry section in WS9 (IO-Link Community vendor table).',
            },
            {
                'q': 'Why does knowing the Device ID matter when ordering a spare sensor?',
                'options': [('A','The master uses it to set the sensor\'s IP address on the IO-Link network'),
                            ('B','It shows the sensor\'s remaining service life in hours'),
                            ('C','It uniquely identifies the exact model, so you order the right part first time and the master can restore parameters automatically after swap')],
                'correct': 'C',
                'note': 'Taught by: the Device Identity section in WS9.',
            },
            {
                'q': 'Port 1 shows PDin hex 5000. Byte 1 bit 4 is set — the instability alarm is active. Correct maintenance response?',
                'options': [('A','Check sensor alignment — the target is likely at the edge of the sensing range'),
                            ('B','No action — an instability alarm means the sensor passed its self-test'),
                            ('C','Replace the sensor immediately — bit 4 means a permanent hardware fault')],
                'correct': 'A',
                'note': 'Taught by: the PDin Bit Table section in WS9 (and WS3).',
            },
            {
                'q': 'Port 2 shows Vendor ID 1586, Device ID 1052673. Which spare should you order?',
                'options': [('A','IFM TV7105 temperature sensor'),
                            ('B','OMRON E2E-X16MB1T12 proximity sensor'),
                            ('C','RS Pro M18 capacitive sensor (Carlo Gavazzi OEM, product code 2377240)')],
                'correct': 'C',
                'note': 'Taught by: the Vendor/Device Registry section in WS9.',
            },
        ],
    },
    {
        'title': 'WS10 — PT100 Temperature Sensors',
        'tables': [
            {
                'header': 'Hex Decode Exercises — decode each PDin hex value to a temperature.',
                'cols': ['PDin Hex', 'Calculation', 'Your Answer'],
                'rows': [
                    ('01 2C FF 00', '0x012C = ?  ÷ 10', '___________'),
                    ('01 96 FF 00', '0x0196 = ?  ÷ 10', '___________'),
                    ('FF 9C FF 00', '0xFF9C = ?  (int16 two\'s complement)  ÷ 10', '___________'),
                    ('FE 70 00 00', '0xFE70 = ?  (int16)  ÷ 10', '___________'),
                ],
                'answers': [
                    ('01 2C FF 00', '0x012C = 300  ÷ 10', '30.0 °C'),
                    ('01 96 FF 00', '0x0196 = 406  ÷ 10', '40.6 °C'),
                    ('FF 9C FF 00', '0xFF9C = −100  ÷ 10', '−10.0 °C'),
                    ('FE 70 00 00', '0xFE70 = −400  ÷ 10', '−40.0 °C  (open-circuit default)'),
                ],
            }
        ],
    },
    {
        'title': 'WS11 — CL50 Pro PDout Encoding Deep-Dive',
        'tables': [
            {
                'header': 'Decode Exercises — interpret each PDout hex value.',
                'cols': ['PDout Hex', 'Your Interpretation'],
                'rows': [
                    ('000100', '___________'),
                    ('004312', '___________'),
                    ('198205', '___________'),
                ],
                'answers': [
                    ('000100', 'Green · Steady · Medium speed · High intensity — AL1350 startup default'),
                    ('004312', 'Orange/Red · Two Colour Flash · Fast · High intensity'),
                    ('198205', 'Lime Green · Flashing · Slow · Low intensity'),
                ],
            },
            {
                'header': 'Encode Exercises — build the PDout hex for each specification.',
                'cols': ['Specification', 'Your Hex'],
                'rows': [
                    ('Blue · Steady · Medium speed · High intensity', '___________'),
                    ('Cyan/Magenta · Two Colour Flash · Fast · High intensity', '___________'),
                ],
                'answers': [
                    ('Blue · Steady · Medium speed · High intensity', '180109'),
                    ('Cyan/Magenta · Two Colour Flash · Fast · High intensity', '0043B7'),
                ],
            }
        ],
    },
]


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT BUILDER
# ─────────────────────────────────────────────────────────────────────────────

GREEN  = RGBColor(0x16, 0xA3, 0x4A)   # green-600
RED    = RGBColor(0xDC, 0x26, 0x26)   # red-600
GRAY   = RGBColor(0x6B, 0x72, 0x80)   # gray-500
NAVY   = RGBColor(0x1E, 0x3A, 0x5F)   # Matrix navy
ORANGE = RGBColor(0xEA, 0x58, 0x0C)   # orange-600


def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = NAVY
    return p


def add_para(doc, text, bold=False, italic=False, color=None, size=None, left_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if left_indent:
        p.paragraph_format.left_indent = left_indent
    return p


def add_option(doc, letter, text, is_correct=False, show_correct=False):
    """Add a single MCQ option line."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)

    if show_correct and is_correct:
        # Bold green tick
        run_tick = p.add_run('✓  ')
        run_tick.bold = True
        run_tick.font.color.rgb = GREEN
        run_letter = p.add_run(f'{letter})  ')
        run_letter.bold = True
        run_letter.font.color.rgb = GREEN
        run_text = p.add_run(text)
        run_text.bold = True
        run_text.font.color.rgb = GREEN
    else:
        run_letter = p.add_run(f'     {letter})  ')
        run_letter.font.color.rgb = GRAY
        run_text = p.add_run(text)
        if show_correct:
            run_text.font.color.rgb = GRAY


def add_table(doc, table_data, show_answers=False):
    cols = table_data['cols']
    rows = table_data['rows'] if not show_answers else table_data['answers']

    p = doc.add_paragraph()
    p.add_run(table_data['header']).bold = True
    p.runs[0].font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(4)

    t = doc.add_table(rows=1 + len(rows), cols=len(cols))
    t.style = 'Table Grid'

    # Header row
    hdr = t.rows[0]
    for i, col_name in enumerate(cols):
        cell = hdr.cells[i]
        cell.text = col_name
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, '1E3A5F')

    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            if show_answers and ci == len(cols) - 1:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = GREEN
            if ri % 2 == 1:
                set_cell_bg(cell, 'F1F5F9')

    doc.add_paragraph()  # spacer


def add_freetext(doc, item, show_answers=False):
    p = doc.add_paragraph()
    p.add_run(item['q']).bold = True

    if show_answers:
        ans_p = doc.add_paragraph()
        ans_p.paragraph_format.left_indent = Cm(0.8)
        run = ans_p.add_run(f'Answer: {item["answer"]}')
        run.bold = True
        run.font.color.rgb = GREEN
    else:
        doc.add_paragraph()  # blank answer space
        doc.add_paragraph()


def add_scenario(doc, scenario, show_answers=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(scenario['title'])
    run.bold = True
    run.font.color.rgb = ORANGE

    if scenario.get('intro'):
        ip = doc.add_paragraph(scenario['intro'])
        ip.paragraph_format.left_indent = Cm(0.4)
        ip.runs[0].italic = True

    if 'options' not in scenario:
        # randomised scenario — just note
        np = doc.add_paragraph()
        np.paragraph_format.left_indent = Cm(0.4)
        run = np.add_run(scenario.get('note', ''))
        run.font.color.rgb = GRAY
        run.italic = True
        return

    for letter, text in scenario['options']:
        is_correct = (letter == scenario.get('correct'))
        add_option(doc, letter, text, is_correct=is_correct, show_correct=show_answers)

    if show_answers and scenario.get('note'):
        np = doc.add_paragraph()
        np.paragraph_format.left_indent = Cm(0.8)
        run = np.add_run(f'Note: {scenario["note"]}')
        run.font.color.rgb = GRAY
        run.italic = True
        run.font.size = Pt(9)


def build_doc(worksheets, course_name, show_answers=False, subtitle=''):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title
    title_p = doc.add_heading(course_name, 0)
    title_p.runs[0].font.color.rgb = NAVY
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph(subtitle)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.runs[0].font.color.rgb = GRAY
    sub_p.runs[0].font.size = Pt(11)
    sub_p.runs[0].italic = True

    if show_answers:
        note_p = doc.add_paragraph('Correct answers shown in green with ✓')
        note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_p.runs[0].font.color.rgb = GREEN
        note_p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    for ws in worksheets:
        add_heading(doc, ws['title'], level=1)

        if ws.get('intro'):
            ip = doc.add_paragraph(ws['intro'])
            ip.paragraph_format.left_indent = Cm(0.4)
            ip.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            ip.runs[0].font.size = Pt(9.5)
            doc.add_paragraph()

        # Tables
        for tbl in ws.get('tables', []):
            add_table(doc, tbl, show_answers=show_answers)

        # Free-text questions
        for ft in ws.get('freetext', []):
            add_freetext(doc, ft, show_answers=show_answers)

        # MCQs
        for qi, q in enumerate(ws.get('questions', []), 1):
            # Question text
            qp = doc.add_paragraph()
            qp.paragraph_format.space_before = Pt(8)
            qp.paragraph_format.space_after  = Pt(3)
            run_num = qp.add_run(f'Q{qi}.  ')
            run_num.bold = True
            qp.add_run(q['q'])

            # Options
            for letter, text in q['options']:
                is_correct = (letter == q['correct'])
                add_option(doc, letter, text, is_correct=is_correct, show_correct=show_answers)

            # Source note (answers doc only)
            if show_answers and q.get('note'):
                np = doc.add_paragraph()
                np.paragraph_format.left_indent = Cm(0.8)
                np.paragraph_format.space_after = Pt(2)
                run = np.add_run(q['note'])
                run.font.color.rgb = GRAY
                run.italic = True
                run.font.size = Pt(9)

        # Scenarios
        for sc in ws.get('scenarios', []):
            add_scenario(doc, sc, show_answers=show_answers)

        doc.add_paragraph()

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# GENERATE FILES
# ─────────────────────────────────────────────────────────────────────────────

files = [
    ('CP0001.docx',         CP0001, 'CP0001 — Interactive Worksheets',   'Student Question Sheet',        False),
    ('CP0001 Answers.docx', CP0001, 'CP0001 — Interactive Worksheets',   'Trainer Answer Key',            True),
    ('CP0002.docx',         CP0002, 'CP0002 — Engineering Worksheets',   'Student Question Sheet',        False),
    ('CP0002 Answers.docx', CP0002, 'CP0002 — Engineering Worksheets',   'Trainer Answer Key',            True),
]

for filename, data, course, subtitle, answers in files:
    path = os.path.join(OUT_DIR, filename)
    doc = build_doc(data, course, show_answers=answers, subtitle=subtitle)
    doc.save(path)
    print(f'Saved: {path}')

print('Done.')
